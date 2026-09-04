import pytesseract
from PIL import Image
import ollama
import json
import os
from datetime import datetime
import pandas as pd
from pytesseract import Output
from pdf2image import convert_from_path


# Configuration Tesseract (adapte si besoin sous Windows)
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# Sur Windows, indique le chemin vers le dossier "bin" de Poppler
# ⚠️ ADAPTE CE CHEMIN avec le numéro de version réel de ton dossier décompressé
POPPLER_PATH = r"C:\Users\POSTE\AppData\Local\spyder-6\Library\bin"


PROMPT_EXTRACTION = """Tu es un système d'extraction de données spécialisé dans les offres commerciales (achat ou vente de matériel).

Voici le texte brut extrait par OCR d'une offre commerciale. Attention : le texte peut être désordonné à cause d'erreurs d'OCR sur les tableaux (colonnes mal alignées, valeurs déplacées). Utilise le contexte pour reconstituer correctement chaque ligne d'article.

---
{texte_ocr}
---

Extrais les informations et retourne UNIQUEMENT un objet JSON valide, sans aucun texte avant ou après, avec exactement cette structure :

{{
  "numero_offre": "",
  "date": "",
  "client": "",
  "fournisseur": "",
  "type": "",
  "articles": [
    {{
      "reference": "",
      "designation": "",
      "quantite": "",
      "prix_unitaire": "",
      "prix_total": ""
    }}
  ],
  "devise": "",
  "delai_livraison": "",
  "conditions_paiement": "",
  "incoterm": "",
  "validite_offre": "",
  "coordonnees": ""
}}

Règles :
- Le tableau contient plusieurs lignes d'articles : crée un objet dans "articles" pour CHAQUE ligne distincte (identifiée par une référence unique, ex. REF-XXXX).
- Pour chaque article, associe correctement la quantité, le prix unitaire et le prix total qui LUI appartiennent, même si l'ordre dans le texte est désorganisé. Utilise la cohérence arithmétique (quantité × prix unitaire ≈ prix total) pour vérifier et corriger les associations si besoin.
- Si une information n'est pas présente dans le texte, laisse la valeur vide "".
- Ne déduis jamais une information qui n'est pas explicitement écrite dans le texte.
- Le champ "type" doit valoir "achat" ou "vente" si déductible du contexte (ex. "OFFRE DE VENTE" ou "OFFRE D'ACHAT" dans le titre), sinon laisse vide.
- Retourne uniquement le JSON, rien d'autre.
"""


def convertir_pdf_en_images(chemin_pdf: str) -> list:
    """Convertit chaque page d'un PDF en image, pour ensuite les passer à l'OCR."""
    images = convert_from_path(chemin_pdf, poppler_path=POPPLER_PATH)
    return images


def extraire_texte_ocr(chemin_image: str) -> str:
    """Version simple (non utilisée dans le pipeline actuel, gardée pour comparaison)."""
    image = Image.open(chemin_image)
    texte = pytesseract.image_to_string(image, lang="fra")
    return texte


def extraire_texte_ocr_structure(image) -> str:
    """
    Version robuste de l'OCR : reconstruit le texte ligne par ligne
    en respectant la position spatiale des mots (meilleure gestion des tableaux).
    Accepte directement un objet image PIL (pas un chemin de fichier).
    """
    data = pytesseract.image_to_data(image, lang="fra", output_type=Output.DICT)
    df = pd.DataFrame(data)

    df = df[df.conf.astype(float) > 0]
    df = df[df.text.str.strip() != ""]

    df["ligne_id"] = df["block_num"].astype(str) + "_" + df["par_num"].astype(str) + "_" + df["line_num"].astype(str)

    lignes_texte = []
    for ligne_id, groupe in df.groupby("ligne_id", sort=False):
        groupe_trie = groupe.sort_values("left")
        ligne = " ".join(groupe_trie["text"].tolist())
        lignes_texte.append((groupe_trie["top"].min(), ligne))

    lignes_texte.sort(key=lambda x: x[0])
    texte_final = "\n".join(ligne for _, ligne in lignes_texte)
    return texte_final
from docx import Document as DocxDocument

def extraire_texte_docx(chemin_docx: str) -> str:
    """Extrait le texte brut d'un fichier Word (.docx) — pas besoin d'OCR, le texte est déjà numérique."""
    doc = DocxDocument(chemin_docx)
    paragraphes = [p.text for p in doc.paragraphs if p.text.strip()]

    # Inclure aussi le texte des tableaux Word, s'il y en a
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphes.append(cell.text.strip())

    return "\n".join(paragraphes)

def extraire_donnees_structurees(texte_ocr: str) -> dict:
    """Envoie le texte OCR à Mistral et récupère un JSON structuré."""
    prompt = PROMPT_EXTRACTION.format(texte_ocr=texte_ocr)

    reponse = ollama.chat(
        model="mistral",
        messages=[{"role": "user", "content": prompt}]
    )

    contenu = reponse["message"]["content"]

    try:
        donnees = json.loads(contenu)
        return donnees
    except json.JSONDecodeError as e:
        print(f"❌ Erreur de parsing JSON : {e}")
        print("Réponse brute du modèle :")
        print(contenu)
        return None


def valider_articles(donnees: dict) -> dict:
    """Vérifie la cohérence quantité × prix_unitaire ≈ prix_total pour chaque article."""
    for article in donnees.get("articles", []):
        try:
            q = float(str(article["quantite"]).replace(",", "."))
            pu = float(str(article["prix_unitaire"]).replace("TND", "").replace(",", ".").strip())
            pt = float(str(article["prix_total"]).replace("TND", "").replace(",", ".").strip())

            if abs(q * pu - pt) > 1:  # tolérance de 1 unité
                article["_alerte"] = "Incohérence détectée : vérifier quantité/prix"
        except (ValueError, KeyError):
            article["_alerte"] = "Champs manquants ou non numériques"
    return donnees


def sauvegarder_json(donnees: dict, dossier_sortie: str = "resultats") -> str:
    """Sauvegarde les données extraites dans un fichier JSON."""
    os.makedirs(dossier_sortie, exist_ok=True)

    numero = donnees.get("numero_offre", "").strip()
    if numero:
        nom_fichier = f"offre_{numero}.json"
    else:
        horodatage = datetime.now().strftime("%Y%m%d_%H%M%S")
        nom_fichier = f"offre_{horodatage}.json"

    chemin_fichier = os.path.join(dossier_sortie, nom_fichier)

    with open(chemin_fichier, "w", encoding="utf-8") as f:
        json.dump(donnees, f, indent=2, ensure_ascii=False)

    return chemin_fichier


def traiter_offre(chemin_fichier_source: str):
    """Pipeline complet : fichier (image ou PDF) -> texte OCR -> JSON -> fichier sauvegardé."""
    print(f"📄 Traitement de : {chemin_fichier_source}")

    extension = os.path.splitext(chemin_fichier_source)[1].lower()

    if extension == ".pdf":
        pages = convertir_pdf_en_images(chemin_fichier_source)
        print(f"📑 PDF converti en {len(pages)} page(s)")
        image = pages[0]  # MVP : on traite uniquement la première page
    else:
        image = Image.open(chemin_fichier_source)

    texte_ocr = extraire_texte_ocr_structure(image)
    print("\n--- Texte extrait par OCR ---")
    print(texte_ocr)

    donnees = extraire_donnees_structurees(texte_ocr)
    if donnees is None:
        print("⚠️ Échec de l'extraction structurée, arrêt du traitement.")
        return

    donnees = valider_articles(donnees)

    print("\n--- Données structurées ---")
    print(json.dumps(donnees, indent=2, ensure_ascii=False))

    chemin_fichier = sauvegarder_json(donnees)
    print(f"\n💾 Résultat sauvegardé : {chemin_fichier}")

def traiter_offre_interface(chemin_fichier_source: str):
    """Version adaptée à Streamlit/scripts externes : gère PDF, image, et Word (.docx)."""
    extension = os.path.splitext(chemin_fichier_source)[1].lower()

    if extension == ".pdf":
        pages = convertir_pdf_en_images(chemin_fichier_source)
        image = pages[0]
        texte_ocr = extraire_texte_ocr_structure(image)
    elif extension == ".docx":
        texte_ocr = extraire_texte_docx(chemin_fichier_source)  # pas d'OCR nécessaire
    else:
        image = Image.open(chemin_fichier_source)
        texte_ocr = extraire_texte_ocr_structure(image)

    donnees = extraire_donnees_structurees(texte_ocr)

    if donnees is None:
        return texte_ocr, None

    donnees = valider_articles(donnees)
    sauvegarder_json(donnees)

    return texte_ocr, donnees